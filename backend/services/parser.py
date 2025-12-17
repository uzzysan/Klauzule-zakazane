"""Document parsing service for PDF and DOCX files."""
from pathlib import Path
from typing import Dict, List, Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from services.ocr import OCRResult, ocr_service


class DocumentSection:
    """Represents a section of a document."""

    def __init__(
        self,
        title: Optional[str],
        content: str,
        start_position: int,
        end_position: int,
        page_number: Optional[int] = None,
    ):
        self.title = title
        self.content = content
        self.start_position = start_position
        self.end_position = end_position
        self.page_number = page_number


class ParsedDocument:
    """Container for parsed document data."""

    def __init__(
        self,
        full_text: str,
        pages: int,
        sections: List[DocumentSection],
        metadata: Dict,
        word_count: int,
        ocr_result: Optional[OCRResult] = None,
    ):
        self.full_text = full_text
        self.pages = pages
        self.sections = sections
        self.metadata = metadata
        self.word_count = word_count
        self.ocr_result = ocr_result


class DocumentParser:
    """Service for parsing PDF and DOCX documents."""

    def parse_pdf(self, file_path: str, language: str = "pol") -> ParsedDocument:
        """
        Parse PDF document using PyMuPDF (fitz).

        Strategy:
        1. Try to extract native text layer
        2. If no text or low quality, use OCR
        3. Extract metadata and structure
        
        PyMuPDF is faster and better preserves document structure than pdfplumber.
        """
        try:
            doc = fitz.open(file_path)
            pages_count = len(doc)

            # Extract metadata
            metadata = {
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
                "creator": doc.metadata.get("creator", ""),
                "producer": doc.metadata.get("producer", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
            }

            # Try native text extraction
            all_text = []
            sections = []

            for page_num in range(pages_count):
                page = doc[page_num]
                page_text = page.get_text("text")  # Extract plain text

                if page_text:
                    all_text.append(page_text)

                    # Create section per page
                    section = DocumentSection(
                        title=f"Page {page_num + 1}",
                        content=page_text,
                        start_position=len("\n\n".join(all_text[:-1])),
                        end_position=len("\n\n".join(all_text)),
                        page_number=page_num + 1,
                    )
                    sections.append(section)

            full_text = "\n\n".join(all_text)

            # Check if we got meaningful text
            ocr_result = None
            if len(full_text.strip()) < 100:
                # Low quality or no text - use OCR
                doc.close()  # Close before OCR
                ocr_result = ocr_service.extract_from_pdf_pages(file_path, language)
                full_text = ocr_result.text

                # Recreate sections from OCR text
                sections = self._create_sections_from_text(full_text, pages_count)
            else:
                doc.close()

            # Count words
            word_count = len(full_text.split())

            return ParsedDocument(
                full_text=full_text,
                pages=pages_count,
                sections=sections,
                metadata=metadata,
                word_count=word_count,
                ocr_result=ocr_result,
            )

        except Exception as e:
            # Return empty document on error
            return ParsedDocument(
                full_text="",
                pages=0,
                sections=[],
                metadata={"error": str(e)},
                word_count=0,
            )

    def parse_docx(self, file_path: str) -> ParsedDocument:
        """
        Parse DOCX document.

        Extracts:
        - Full text
        - Paragraphs as sections
        - Metadata
        - Structure (headings, lists)
        """
        try:
            doc = DocxDocument(file_path)

            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }

            # Extract text and structure
            sections = []
            all_text = []
            current_position = 0

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()

                if text:
                    all_text.append(text)

                    # Check if it's a heading
                    is_heading = paragraph.style.name.startswith("Heading")

                    section = DocumentSection(
                        title=text if is_heading else None,
                        content=text,
                        start_position=current_position,
                        end_position=current_position + len(text),
                    )
                    sections.append(section)

                    current_position += len(text) + 2  # +2 for newlines

            full_text = "\n\n".join(all_text)
            word_count = len(full_text.split())

            # Estimate pages (rough: 500 words per page)
            pages = max(1, word_count // 500)

            return ParsedDocument(
                full_text=full_text,
                pages=pages,
                sections=sections,
                metadata=metadata,
                word_count=word_count,
            )

        except Exception as e:
            return ParsedDocument(
                full_text="",
                pages=0,
                sections=[],
                metadata={"error": str(e)},
                word_count=0,
            )

    def parse_image(self, file_path: str, language: str = "pol") -> ParsedDocument:
        """
        Parse image file using OCR.
        """
        try:
            # Read image file
            with open(file_path, "rb") as f:
                image_data = f.read()

            # Run OCR
            ocr_result = ocr_service.extract_text_from_image(
                image_data,
                language=language,
                preprocess=True,
            )

            # Create single section
            sections = [
                DocumentSection(
                    title="Image Content",
                    content=ocr_result.text,
                    start_position=0,
                    end_position=len(ocr_result.text),
                    page_number=1,
                )
            ]

            word_count = len(ocr_result.text.split())

            return ParsedDocument(
                full_text=ocr_result.text,
                pages=1,
                sections=sections,
                metadata={
                    "source": "image_ocr",
                    "confidence": ocr_result.confidence,
                },
                word_count=word_count,
                ocr_result=ocr_result,
            )

        except Exception as e:
            return ParsedDocument(
                full_text="",
                pages=0,
                sections=[],
                metadata={"error": str(e)},
                word_count=0,
            )

    def _create_sections_from_text(self, text: str, pages: int) -> List[DocumentSection]:
        """Create sections by splitting text into pages."""
        # Split text evenly across pages
        words = text.split()
        words_per_page = max(1, len(words) // pages)

        sections = []
        for i in range(pages):
            start_idx = i * words_per_page
            end_idx = (i + 1) * words_per_page if i < pages - 1 else len(words)

            page_words = words[start_idx:end_idx]
            page_text = " ".join(page_words)

            section = DocumentSection(
                title=f"Page {i + 1}",
                content=page_text,
                start_position=len(" ".join(words[:start_idx])),
                end_position=len(" ".join(words[:end_idx])),
                page_number=i + 1,
            )
            sections.append(section)

        return sections


# Singleton instance
document_parser = DocumentParser()
